# -*- coding: utf-8 -*-
"""
reader.py — Đọc văn bản pháp luật từ .docx hoặc .pdf ra text sạch.

Giải quyết ba lỗi đã xác định trong knowledge.py hiện tại:
  1. python-docx bỏ qua doc.tables  -> duyệt body theo thứ tự, gồm cả bảng
  2. PDF có header/footer lặp mỗi trang -> phát hiện và lọc
  3. Không phân biệt được scan/text-layer -> báo lỗi rõ thay vì trả chuỗi rỗng
"""

from __future__ import annotations

import re
from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path

MIN_CHARS_PER_PAGE = 200          # dưới ngưỡng này coi như không có text layer
HEADER_MIN_RATIO = 0.6            # dòng xuất hiện >=60% số trang -> header/footer
TABLE_MIN_COLS = 2                # Nx1 là khung kẻ layout, không phải bảng


class NoTextLayer(Exception):
    """PDF là bản scan — phải OCR trước khi ingest."""


@dataclass
class Document:
    path: Path
    text: str
    source_format: str                     # "docx" | "pdf"
    n_pages: int = 0
    n_tables: int = 0
    removed_headers: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)

    @property
    def n_chars(self) -> int:
        return len(self.text)


# ---------------------------------------------------------------- DOCX

def _table_to_text(table) -> str:
    """
    Bảng -> text có phân tách rõ ràng.

    Dùng ' | ' giữa các ô để chunker và LLM còn nhận ra ranh giới cột.
    Nối bằng khoảng trắng sẽ làm các ô dính vào nhau thành câu vô nghĩa.
    """
    lines = []
    for row in table.rows:
        cells = [c.text.strip().replace("\n", " ") for c in row.cells]
        # bỏ ô lặp do merge ngang
        dedup, prev = [], None
        for c in cells:
            if c != prev:
                dedup.append(c)
            prev = c
        if any(dedup):
            lines.append(" | ".join(dedup))
    return "\n".join(lines)


W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _textbox_text(element) -> str:
    """
    Text nằm trong textbox (w:txbxContent).

    File .docx do công cụ OCR/convert sinh ra thường nhét text vào textbox
    để giữ nguyên vị trí trên trang. python-docx bỏ qua hoàn toàn phần này,
    nên một file trông đầy chữ trong Word lại trích ra chuỗi rỗng.
    """
    out = []
    for tb in element.iter(f"{W_NS}txbxContent"):
        for p in tb.iter(f"{W_NS}p"):
            t = "".join(n.text or "" for n in p.iter(f"{W_NS}t")).strip()
            if t:
                out.append(t)
    return "\n".join(out)


def read_docx(path: Path) -> Document:
    import docx
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = docx.Document(str(path))
    body = doc.element.body

    # Duyệt theo THỨ TỰ XUẤT HIỆN trong tài liệu. Lặp doc.paragraphs rồi
    # doc.tables sẽ đẩy toàn bộ bảng xuống cuối, phá vỡ mạch Điều/Khoản.
    parts, n_tables, n_textbox = [], 0, 0
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            t = Paragraph(child, doc).text.strip()
            if t:
                parts.append(t)
            tb = _textbox_text(child)
            if tb:
                parts.append(tb)
                n_textbox += 1
        elif tag == "tbl":
            table = Table(child, doc)
            t = _table_to_text(table)
            if t:
                parts.append(t)
                n_tables += 1
            tb = _textbox_text(child)
            if tb:
                parts.append(tb)
                n_textbox += 1

    text = "\n".join(parts)
    warnings = []

    if n_textbox:
        warnings.append(
            f"{n_textbox} khối text nằm trong textbox — dấu hiệu file được "
            "convert từ PDF. Thứ tự đọc có thể không khớp thứ tự trên trang; "
            "kiểm tra chuỗi Điều bằng verify_ocr.py."
        )

    n_images = len([r for r in doc.part.rels.values() if "image" in r.reltype])
    if len(text) < 1000 and n_images:
        warnings.append(
            f"Chỉ {len(text)} ký tự nhưng có {n_images} ảnh nhúng. "
            "File này là bản scan đã đổi vỏ sang .docx, KHÔNG có text thật. "
            "Cần OCR chứ không phải convert."
        )

    return Document(
        path=path,
        text=text,
        source_format="docx",
        n_tables=n_tables,
        warnings=warnings,
    )


# ---------------------------------------------------------------- PDF

def _find_repeated_lines(pages: list[str]) -> list[str]:
    """Dòng lặp ở đầu/cuối phần lớn các trang -> header/footer."""
    if len(pages) < 3:
        return []
    counter = Counter()
    for text in pages:
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        for line in set(lines[:3] + lines[-3:]):
            if 3 < len(line) < 120:
                counter[line] += 1
    threshold = len(pages) * HEADER_MIN_RATIO
    return [l for l, c in counter.items() if c >= threshold]


def _strip_page_numbers(line: str) -> bool:
    """True nếu dòng chỉ là số trang hoặc dạng '12/34'."""
    s = line.strip()
    return bool(re.fullmatch(r"-?\s*\d{1,4}\s*(/\s*\d{1,4})?\s*-?", s))


def read_pdf(path: Path) -> Document:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = [(p.extract_text() or "") for p in reader.pages]
    n_pages = len(pages)
    total = sum(len(p) for p in pages)

    if total // max(n_pages, 1) < MIN_CHARS_PER_PAGE:
        raise NoTextLayer(
            f"{path.name}: {total // max(n_pages, 1)} ký tự/trang. "
            "Bản scan — chạy OCR (ocrmypdf -l vie) hoặc tìm bản Công báo có text layer."
        )

    headers = _find_repeated_lines(pages)
    header_set = set(headers)

    cleaned = []
    for page in pages:
        for line in page.split("\n"):
            s = line.strip()
            if not s or s in header_set or _strip_page_numbers(s):
                continue
            cleaned.append(s)

    text = "\n".join(cleaned)

    warnings = []
    n_tables = _count_real_tables(path)
    if n_tables:
        warnings.append(
            f"Có {n_tables} bảng thật trong PDF. pypdf làm bẹp bảng — "
            "nếu bảng mang nội dung quy phạm (biểu thuế, danh mục) thì trích riêng "
            "bằng pdfplumber trước khi tin vào text này."
        )

    return Document(
        path=path,
        text=text,
        source_format="pdf",
        n_pages=n_pages,
        n_tables=n_tables,
        removed_headers=headers,
        warnings=warnings,
    )


def _count_real_tables(path: Path) -> int:
    """Đếm bảng >= TABLE_MIN_COLS cột. Bảng Nx1 là khung kẻ layout."""
    try:
        import pdfplumber
    except ImportError:
        return 0
    n = 0
    try:
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                n += sum(1 for t in page.find_tables()
                         if len(t.columns) >= TABLE_MIN_COLS)
    except Exception:
        return 0
    return n


# ---------------------------------------------------------------- chung

def normalize(text: str) -> str:
    """Chuẩn hóa nhẹ, không đụng tới nội dung."""
    text = text.replace("\u00a0", " ").replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def read(path: str | Path) -> Document:
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".docx":
        doc = read_docx(path)
    elif suffix == ".pdf":
        doc = read_pdf(path)
    elif suffix == ".doc":
        raise ValueError(
            f"{path.name}: .doc cũ không đọc trực tiếp được. "
            "Convert trước: soffice --headless --convert-to docx"
        )
    else:
        raise ValueError(f"{path.name}: không hỗ trợ đuôi {suffix}")

    doc.text = normalize(doc.text)
    return doc